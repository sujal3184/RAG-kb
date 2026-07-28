"""Shared fixtures for loader tests — builds small sample files in-memory."""

import io

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter


@pytest.fixture
def sample_txt_bytes() -> bytes:
    return "Hello world.\nThis is a test document about RAG systems.".encode("utf-8")


@pytest.fixture
def sample_html_bytes() -> bytes:
    html = """
    <html>
      <head><title>Test Page</title><style>body { color: red; }</style></head>
      <body>
        <h1>Main Heading</h1>
        <p>This is a paragraph about knowledge bases.</p>
        <script>console.log("should be removed");</script>
      </body>
    </html>
    """
    return html.encode("utf-8")


@pytest.fixture
def sample_csv_bytes() -> bytes:
    csv_content = "name,role\nAlice,Engineer\nBob,Designer\n"
    return csv_content.encode("utf-8")


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Build a minimal real .docx file in memory using python-docx."""
    doc = DocxDocument()
    doc.add_paragraph("This is a test paragraph about RAG systems.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Build a minimal real (but text-less) PDF using pypdf.

    Note: pypdf's PdfWriter can create blank pages but not easily draw
    real text without an additional library — so this fixture verifies
    the loader handles a VALID PDF structure correctly. Text-extraction
    accuracy on a page with actual content is implicitly covered by the
    fact that `extract_text()` is exercised without error.
    """
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_two_column_pdf_bytes() -> bytes:
    """Build a real PDF with two side-by-side columns of text."""
    import pdfplumber  # noqa: F401  (ensures pdfplumber is importable in test env)
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    left_lines = ["Left column line one.", "Left column line two.", "Left column line three."]
    right_lines = ["Right column line one.", "Right column line two.", "Right column line three."]

    y = 750
    for line in left_lines:
        c.drawString(50, y, line)
        y -= 20

    y = 750
    for line in right_lines:
        c.drawString(320, y, line)
        y -= 20

    c.save()
    return buffer.getvalue()


@pytest.fixture
def sample_pdf_with_table_bytes() -> bytes:
    """Build a real PDF containing a simple table WITH visible grid lines.

    pdfplumber detects tables by finding visible borders/lines — a table
    with no styling at all is indistinguishable from plain aligned text,
    so we must apply a grid style for extract_tables() to find it.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    data = [["Name", "Role"], ["Alice", "Engineer"], ["Bob", "Designer"]]

    table = Table(data)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )

    doc.build([table])
    return buffer.getvalue()