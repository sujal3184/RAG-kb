"""Loader for Word documents (.docx)."""

import io
import zipfile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.loaders.base import DocumentLoader, LoadedDocument
from app.loaders.exceptions import LoaderError


class DocxLoader(DocumentLoader):
    """Loads .docx files, extracting text from paragraphs and tables."""

    def load(self, content: bytes) -> LoadedDocument:
        """Extract all paragraph and table text from a .docx file.

        Raises:
            LoaderError: if the file isn't a valid .docx (e.g. it's
                actually a legacy .doc file, corrupted, or not a Word
                document at all).
        """
        try:
            docx_document = DocxDocument(io.BytesIO(content))
        except (PackageNotFoundError, zipfile.BadZipFile, KeyError) as exc:
            raise LoaderError(
                "Could not open .docx file — it may be corrupted, empty, or a "
                "legacy .doc file (only modern .docx is supported)"
            ) from exc

        paragraph_texts = [p.text for p in docx_document.paragraphs if p.text.strip()]

        table_texts: list[str] = []
        for table in docx_document.tables:
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]
                if any(cell_texts):
                    table_texts.append(" | ".join(cell_texts))

        full_text = "\n".join(paragraph_texts + table_texts)

        return LoadedDocument(
            text=full_text,
            metadata={
                "paragraph_count": len(paragraph_texts),
                "table_count": len(docx_document.tables),
            },
        )